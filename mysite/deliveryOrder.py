#!/usr/bin/env python

import tools

import os
import time
from copy import deepcopy

from docx import Document

from flask import send_from_directory



def deliveryorder(ID):


    sql = "SELECT order_id, ccn ,client , Orders.phone,cargo_type , Warehouse.company , Warehouse.address,   Orders.address ,amount , weight,volume,delivery_type, User.name \
           FROM Orders JOIN Warehouse ON Warehouse.warehouse_id = Orders.warehouse  JOIN User ON User.user_id = Orders.operator WHERE order_id = '%s' " % ID

    results = tools.queryMysql(sql)
    if not results:
        return "没有相关数据，请核准单号之后，再进行相关操作！"

    result = []
    for item in results[0]:
        if item:
            result.append(item)
        else:
            result.append('')

    DO = Document('/home/KTRANS/mysite/static/Delivery_Order_Template.docx')


    filename = DO.tables[0].cell(0,0).paragraphs[0]
    filename.text = 'File No.' + ID
    DO.tables[0].cell(0,0).add_paragraph('CCN#' + result[1])

    consignee = DO.tables[0].cell(0,1).paragraphs[0]
    consignee.text = 'Consignee：' + result[2]
    DO.tables[0].cell(0,1).add_paragraph('Phone number：' + result[3])

    today = time.strftime("%Y-%m-%d",time.localtime(time.time()))
    DO.tables[0].cell(1,0).add_paragraph('Issue Date:%s operator:  %s'%(today,result[12] ))
    DO.tables[0].cell(1,1).add_paragraph('Goods description: %s '%result[4] )


    DO.tables[0].cell(3,0).add_paragraph('From:   %s warehouse at: %s'%(result[5],result[6]))

    DO.tables[0].cell(4,0).add_paragraph('To:   %s '%result[7])

    DO.tables[0].cell(5,0).add_paragraph('Piece:  %s '%result[8])
    DO.tables[0].cell(5,1).add_paragraph('Weight：  %s '%result[9])
    DO.tables[0].cell(5,3).add_paragraph('Dimension:  %s '%result[10])

    DO.tables[1].cell(1,1).add_paragraph('REMARK:  %s '%result[11])


    DO.save('/home/KTRANS/mysite/static/files/deliveryOrder/Delivery_Order_%s.docx'%(ID))


def FBA_deliveryorder(warehouse, cargo_id, pallets, date, cargo_list):
    """生成 FBA Delivery Order Word 文档，安全处理空单元格和多行填充"""


    # 打开模板
    DO = Document('/home/KTRANS/mysite/static/FBA_deliveryorder_Template.docx')

    # 获取仓库地址
    sql = "SELECT address FROM FBA WHERE name = '%s'" % warehouse

    address_tuple = tools.queryMysql(sql)[0]  # 可能是 ('123 Street, City',)
    address = str(address_tuple[0])

    DO.tables[0].cell(1,0).add_paragraph(warehouse + ' ' + address)
    DO.tables[0].cell(1,1).add_paragraph(warehouse)
    DO.tables[0].cell(1,3).add_paragraph(pallets)


    # 填充表格1
    table = DO.tables[1]
    template_row = table.rows[1]

    # 删除多余行，保留模板行
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[1]._tr)

    for data in cargo_list:
        new_row = deepcopy(template_row._tr)
        table._tbl.append(new_row)
        cells = table.rows[-1].cells

        for i, value in enumerate(data):
            cell = cells[i]

            para = cell.paragraphs[0]  # 直接取模板行的第一个段落
            if not para.runs:
                para.add_run()
            para.runs[0].text = str(value)

            # 清空多余 run
            for run in para.runs[1:]:
                run.text = ""

            # 删除其他多余段落
            for p in cell.paragraphs[1:]:
                p.clear()  # 或者 cell._element.remove(p._p)


    filePath = '/home/KTRANS/mysite/static/files/FBA/'

    filename = '%s--%s-RonnieFile%spallets%s.docx' % (
        warehouse, cargo_id, pallets, date
    )

    full_path = os.path.join(filePath, filename)

    DO.save(full_path)

    return send_from_directory(filePath, filename, as_attachment=True)




